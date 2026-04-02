"""readers/docx.py — Word (.docx) → Markdown変換（python-docx使用）"""

from __future__ import annotations


def read(filepath: str) -> str:
    """Word文書をMarkdown形式に変換して返す。"""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("Wordの読み込みにはpython-docxが必要です。\npip install python-docx")

    doc = Document(filepath)
    lines = []

    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            lines.append("")
            continue

        if style.startswith("Heading 1"):
            lines.append(f"# {text}")
        elif style.startswith("Heading 2"):
            lines.append(f"## {text}")
        elif style.startswith("Heading 3"):
            lines.append(f"### {text}")
        elif style.startswith("Heading"):
            lines.append(f"#### {text}")
        elif style.startswith("List"):
            lines.append(f"- {text}")
        else:
            lines.append(text)

    # テーブル
    for table in doc.tables:
        lines.append("")
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
            if i == 0:
                lines.append("|" + "|".join(["---"] * len(cells)) + "|")
        lines.append("")

    return "\n".join(lines)
